# use "python -m IPython" for testing
import json
import docx

with open('gpaasch_2020_resume.json') as f:
    resume = json.loads(f.read())

length_in_points = docx.shared.Pt
document = docx.Document()
align_center = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
align_left = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

email_phone_website_address = \
    resume['basics']['email'] + ' | ' + \
    resume['basics']['phone'] + ' | ' + \
    resume['basics']['website'] + ' | ' + \
    resume['basics']['location']['address'] + '\n'

length_of_dashes = len(email_phone_website_address) - 1

contact_info = document.add_paragraph()
contact_info.alignment = align_center
contact_info.add_run(resume['basics']['name'] + '\n', 'Title Char')
contact_info.add_run(email_phone_website_address)
contact_info.add_run('- ' * length_of_dashes)
contact_info.paragraph_format.space_after = length_in_points(0)

certification = document.add_paragraph()
certification.alignment = align_left
for award in resume['awards']:
    certification.add_run(award['title'] + '\n')
certification.add_run('- ' * length_of_dashes)
certification.paragraph_format.space_before = length_in_points(0)

document.save('gpaasch_2020_resume.docx')